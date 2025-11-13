# actas_routes.py — generación de ACTAS (DOCX + PDF) para Mediazion
import os
from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from datetime import datetime
from io import BytesIO
import requests

# DOCX
from docx import Document
from docx.shared import Inches

# PDF
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

actas_router = APIRouter()

UPLOAD_DIR = "uploads/actas"
os.makedirs(UPLOAD_DIR, exist_ok=True)

class ActaIn(BaseModel):
    case_no: str
    date_iso: str
    mediator_alias: str = ""
    parties: str = ""
    summary: str
    agreements: str = ""
    confidentiality: bool = True
    logo_url: str = None


# ---------------------------------------------------------------------
# ---------------              DOCX             ------------------------
# ---------------------------------------------------------------------
@actas_router.post("/actas/render_docx")
def render_docx(data: ActaIn):

    try:
        doc = Document()

        # Logo
        if data.logo_url:
            try:
                img_bytes = requests.get(data.logo_url, timeout=5).content
                doc.add_picture(BytesIO(img_bytes), width=Inches(1.5))
            except:
                pass  # no romper si no hay logo

        doc.add_heading("ACTA DE MEDIACIÓN", level=1)
        doc.add_paragraph(f"Expediente: {data.case_no}")
        doc.add_paragraph(f"Fecha: {data.date_iso}")

        if data.mediator_alias:
            doc.add_paragraph(f"Mediador/a: {data.mediator_alias}")

        if data.parties:
            doc.add_paragraph(f"Partes: {data.parties}")

        doc.add_heading("Resumen / Contenido", level=2)
        doc.add_paragraph(data.summary)

        if data.agreements:
            doc.add_heading("Acuerdos", level=2)
            doc.add_paragraph(data.agreements)

        if data.confidentiality:
            doc.add_paragraph("")
            doc.add_paragraph(
                "Cláusula de confidencialidad:\n"
                "Las partes se comprometen a mantener la confidencialidad del proceso."
            )

        # Guardar DOCX
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        fname = f"acta_{data.case_no.replace(' ','_')}_{timestamp}.docx"
        fpath = os.path.join(UPLOAD_DIR, fname)
        doc.save(fpath)

        return {"ok": True, "url": f"/uploads/actas/{fname}"}

    except Exception as e:
        raise HTTPException(500, f"Error generando DOCX: {str(e)}")


# ---------------------------------------------------------------------
# ----------------             PDF              ------------------------
# ---------------------------------------------------------------------
@actas_router.post("/actas/render_pdf")
def render_pdf(data: ActaIn):

    try:
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        fname = f"acta_{data.case_no.replace(' ','_')}_{timestamp}.pdf"
        fpath = os.path.join(UPLOAD_DIR, fname)

        c = canvas.Canvas(fpath, pagesize=A4)
        width, height = A4
        y = height - 70

        # Logo
        if data.logo_url:
            try:
                img_bytes = BytesIO(requests.get(data.logo_url, timeout=5).content)
                c.drawImage(img_bytes, 40, y - 40, width=100, preserveAspectRatio=True, mask='auto')
                y -= 60
            except:
                pass

        # Cabecera
        c.setFont("Helvetica-Bold", 16)
        c.drawString(150, y, "ACTA DE MEDIACIÓN")
        y -= 40

        c.setFont("Helvetica", 11)
        lines = [
            f"Expediente: {data.case_no}",
            f"Fecha: {data.date_iso}",
        ]

        if data.mediator_alias:
            lines.append(f"Mediador/a: {data.mediator_alias}")
        if data.parties:
            lines.append(f"Partes: {data.parties}")

        lines.append("")
        lines.append("Resumen / Contenido:")
        for l in data.summary.split("\n"):
            lines.append(l)

        if data.agreements:
            lines.append("")
            lines.append("Acuerdos:")
            for l in data.agreements.split("\n"):
                lines.append(l)

        if data.confidentiality:
            lines.append("")
            lines.append(
                "Cláusula de confidencialidad:\n"
                "Las partes se comprometen a mantener la confidencialidad del proceso."
            )

        # Escribir líneas
        for line in lines:
            for segment in line.split("\n"):
                c.drawString(40, y, segment)
                y -= 18
                if y < 60:
                    c.showPage()
                    c.setFont("Helvetica", 11)
                    y = height - 70

        c.save()

        return {"ok": True, "url": f"/uploads/actas/{fname}"}

    except Exception as e:
        raise HTTPException(500, f"Error generando PDF: {str(e)}")
