# actas_docx_logo.py — Renderizar actas DOCX con logo en cabecera y listarlas por caso
from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import Optional
from uuid import uuid4
from pathlib import Path
import tempfile
import datetime as dt

import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests

router = APIRouter(prefix="/api/actas", tags=["actas"])


class ActaPayload(BaseModel):
    case_no: str
    date_iso: str
    mediator_alias: str
    parties: str
    summary: str
    agreements: str
    confidentiality: bool = True
    location: Optional[str] = "España"
    logo_url: Optional[str] = None
    logo_mode: Optional[str] = "normal"
    logo_width_cm: float = 9.0
    caso_id: Optional[str] = None  # vínculo opcional al ID de caso


class ActaInfo(BaseModel):
    filename: str
    url: str
    created_at: Optional[str] = None


BASE_DIR = Path("generated_actas")
BASE_DIR.mkdir(parents=True, exist_ok=True)


@router.post("/render_docx")
def render_docx_acta(body: ActaPayload):
    """Genera un DOCX de acta con cabecera de logo.

    El archivo se guarda en generated_actas y se devuelve una URL descargable.
    Si se facilita caso_id, el nombre del fichero incluye ese ID para poder listarlo después.
    """
    if not body.case_no or not body.date_iso:
        raise HTTPException(400, "Faltan datos básicos del acta (expediente y fecha).")

    doc = docx.Document()

    # Cabecera con logo
    if body.logo_url:
        try:
            resp = requests.get(body.logo_url, timeout=8)
            if resp.ok:
                tmpdir = Path(tempfile.gettempdir())
                img_path = tmpdir / f"logo_acta_{uuid4().hex}.png"
                img_path.write_bytes(resp.content)

                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(img_path), width=Cm(body.logo_width_cm))
        except Exception:
            # No rompemos la generación de acta si hay un problema con el logo
            pass

    # Título
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("ACTA DE MEDIACIÓN")
    run_title.bold = True

    # Datos básicos
    doc.add_paragraph(f"Nº de expediente: {body.case_no}")
    doc.add_paragraph(f"Fecha: {body.date_iso}")
    if body.location:
        doc.add_paragraph(f"Lugar: {body.location}")
    doc.add_paragraph(f"Mediador/a: {body.mediator_alias}")
    doc.add_paragraph(f"Partes intervinientes: {body.parties}")
    doc.add_paragraph("")  # línea en blanco

    # Antecedentes / Hechos
    p1 = doc.add_paragraph()
    p1.add_run("Antecedentes / Hechos:").bold = True
    doc.add_paragraph(body.summary)
    doc.add_paragraph("")  # línea en blanco

    # Acuerdos
    p2 = doc.add_paragraph()
    p2.add_run("Acuerdos y compromisos:").bold = True
    doc.add_paragraph(body.agreements)

    # Confidencialidad
    if body.confidentiality:
        doc.add_paragraph("")
        p3 = doc.add_paragraph()
        p3.add_run("Confidencialidad:").bold = True
        doc.add_paragraph(
            "Las partes manifiestan que han sido informadas del carácter confidencial de la mediación y se comprometen "
            "a no utilizar en otros procedimientos la información generada en este espacio salvo acuerdo expreso."
        )

    # Firmas
    doc.add_paragraph("")
    doc.add_paragraph("Firmas:")
    doc.add_paragraph("La persona mediadora: ________________________________")
    doc.add_paragraph("Las partes intervinientes: ____________________________")

    # Guardar DOCX en disco
    prefix = f"acta_caso-{body.caso_id}_" if body.caso_id else "acta_"
    file_id = uuid4().hex
    filename = f"{prefix}{file_id}.docx"
    file_path = BASE_DIR / filename
    doc.save(str(file_path))

    # URL pública (ajusta a tu servidor de estáticos)
    url = f"/static/actas/{filename}"

    return {"ok": True, "url": url}


@router.get("")
def list_actas(caso_id: str):
    """Lista actas DOCX vinculadas a un caso concreto.

    Busca ficheros en generated_actas cuyo nombre empiece por 'acta_caso-{caso_id}_'.
    """
    if not caso_id:
        raise HTTPException(400, "caso_id requerido")

    prefix = f"acta_caso-{caso_id}_"
    results = []

    if not BASE_DIR.exists():
        return []

    for entry in BASE_DIR.iterdir():
        if not entry.is_file():
            continue
        if not entry.name.startswith(prefix):
            continue
        stat = entry.stat()
        created_at = dt.datetime.fromtimestamp(stat.st_mtime).isoformat()
        url = f"/static/actas/{entry.name}"
        results.append(
            ActaInfo(filename=entry.name, url=url, created_at=created_at).dict()
        )

    # Ordenamos por fecha descendente
    results.sort(key=lambda a: a.get("created_at") or "", reverse=True)
    return results
